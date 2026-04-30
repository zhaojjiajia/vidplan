from __future__ import annotations

import html
import json
import re
from dataclasses import dataclass
from io import BytesIO
from typing import Any

import markdown as markdown_lib
from docx import Document

from apps.plans.models import SeriesPlan, VideoPlan

from .markdown import plan_to_markdown


class ExportRenderError(Exception):
    pass


@dataclass(frozen=True)
class ExportedFile:
    body: bytes
    content_type: str
    filename: str


def render_plan(plan: VideoPlan, fmt: str) -> ExportedFile:
    normalized = normalize_format(fmt)
    filename = f"{safe_filename(plan.title or 'plan')}.{normalized}"

    if normalized == "md":
        return ExportedFile(
            body=plan_to_markdown(plan).encode("utf-8"),
            content_type="text/markdown; charset=utf-8",
            filename=filename,
        )
    if normalized == "pdf":
        return ExportedFile(
            body=plan_to_pdf(plan),
            content_type="application/pdf",
            filename=filename,
        )
    if normalized == "docx":
        return ExportedFile(
            body=plan_to_docx(plan),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=filename,
        )
    raise ValueError(f"暂不支持的导出格式: {fmt}")


def render_series(series: SeriesPlan, fmt: str) -> ExportedFile:
    normalized = normalize_format(fmt)
    filename = f"{safe_filename(series.title or 'series')}.{normalized}"

    if normalized == "md":
        return ExportedFile(
            body=series_to_markdown(series).encode("utf-8"),
            content_type="text/markdown; charset=utf-8",
            filename=filename,
        )
    if normalized == "docx":
        return ExportedFile(
            body=series_to_docx(series),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=filename,
        )
    raise ValueError(f"暂不支持的系列导出格式: {fmt}")


def normalize_format(fmt: str) -> str:
    value = (fmt or "md").lower().strip()
    aliases = {"markdown": "md", "word": "docx"}
    return aliases.get(value, value)


def safe_filename(title: str) -> str:
    safe = re.sub(r"[^\w\u4e00-\u9fa5\-]+", "_", title).strip("_")[:60]
    return safe or "plan"


def series_to_markdown(series: SeriesPlan) -> str:
    lines: list[str] = []
    lines.append(f"# {series.title or '未命名系列'}")
    lines.append("")
    lines.append(f"> {series.summary}" if series.summary else "> _暂无简介_")
    lines.append("")

    lines.append("## 基本信息")
    lines.append("")
    info = [
        ("方向", series.direction),
        ("目标平台", series.target_platform or "—"),
        ("目标观众", series.target_audience or "—"),
        ("更新频率", series.update_frequency or "—"),
        ("单集时长 (秒)", series.episode_duration_seconds),
        ("计划集数", series.planned_episodes),
        ("状态", series.get_status_display()),
        ("最后更新", series.updated_at.strftime("%Y-%m-%d %H:%M")),
    ]
    for key, value in info:
        lines.append(f"- **{key}**: {_stringify(value)}")
    lines.append("")

    lines.extend(_markdown_mapping_section("系列定位", series.positioning or {}))
    lines.extend(_markdown_mapping_section("单集模板", series.episode_template or {}))
    lines.extend(_markdown_mapping_section("视觉风格", series.visual_style or {}))
    lines.extend(_markdown_mapping_section("标题风格", series.title_style or {}))
    lines.extend(_markdown_list_section("初始选题", series.initial_topics or []))
    lines.extend(_markdown_assets_section(series))
    lines.extend(_markdown_episodes_section(series))

    return "\n".join(lines).rstrip() + "\n"


def _markdown_mapping_section(title: str, data: dict[str, Any]) -> list[str]:
    if not data:
        return []
    out = [f"## {title}", ""]
    for key, value in data.items():
        out.append(f"- **{key}**: {_stringify(value)}")
    out.append("")
    return out


def _markdown_list_section(title: str, items: list[Any]) -> list[str]:
    if not items:
        return []
    out = [f"## {title}", ""]
    for index, item in enumerate(items, 1):
        out.append(f"{index}. {_stringify(item)}")
    out.append("")
    return out


def _markdown_assets_section(series: SeriesPlan) -> list[str]:
    groups = [
        ("人物资产", series.characters.all()),
        ("风格资产", series.styles.all()),
        ("世界观资产", series.worldviews.all()),
        ("栏目资产", series.columns.all()),
    ]
    out: list[str] = []
    for title, queryset in groups:
        assets = list(queryset)
        if not assets:
            continue
        out.extend([f"## {title}", ""])
        for asset in assets:
            out.append(f"### {asset.name}")
            if asset.fixed_traits:
                out.append(f"- **固定特征**: {_stringify(asset.fixed_traits)}")
            if asset.payload:
                out.append(f"- **详细设定**: {_stringify(asset.payload)}")
            out.append("")
    return out


def _markdown_episodes_section(series: SeriesPlan) -> list[str]:
    episodes = list(series.episodes.all())
    if not episodes:
        return []
    out = ["## 单集清单", ""]
    out.append("| # | 标题 | 状态 | 时长 | 更新时间 |")
    out.append("|---|---|---|---|---|")
    for index, episode in enumerate(episodes, 1):
        out.append(
            "| "
            + " | ".join(
                [
                    str(index),
                    _md_cell(episode.title or "未命名单集"),
                    _md_cell(episode.get_status_display()),
                    _md_cell(f"{episode.duration_seconds}s"),
                    _md_cell(episode.updated_at.strftime("%Y-%m-%d %H:%M")),
                ]
            )
            + " |"
        )
    out.append("")
    return out


def _md_cell(value: Any) -> str:
    s = _stringify(value)
    return s.replace("|", "\\|").replace("\n", "<br/>")


def plan_to_pdf(plan: VideoPlan) -> bytes:
    markdown_text = plan_to_markdown(plan)
    body_html = markdown_lib.markdown(markdown_text, extensions=["tables", "nl2br"])
    document_html = f"""
<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <style>
    @page {{ size: A4; margin: 18mm 16mm; }}
    body {{
      font-family: "PingFang SC", "Noto Sans CJK SC", "Microsoft YaHei", sans-serif;
      color: #1f2937;
      font-size: 12px;
      line-height: 1.65;
    }}
    h1 {{ font-size: 24px; margin: 0 0 12px; }}
    h2 {{ font-size: 17px; margin: 22px 0 8px; border-bottom: 1px solid #d8dee6; padding-bottom: 4px; }}
    blockquote {{ color: #5b6472; border-left: 3px solid #9aa4b2; margin: 0 0 16px; padding-left: 12px; }}
    table {{ width: 100%; border-collapse: collapse; table-layout: fixed; margin-top: 8px; }}
    th, td {{ border: 1px solid #d8dee6; padding: 6px; vertical-align: top; word-break: break-word; }}
    th {{ background: #f3f6f9; font-weight: 600; }}
    code {{ white-space: pre-wrap; }}
  </style>
</head>
<body>{body_html}</body>
</html>
"""
    return _render_pdf_with_weasyprint(document_html)


def _render_pdf_with_weasyprint(document_html: str) -> bytes:
    try:
        from weasyprint import HTML
    except (ImportError, OSError) as exc:
        raise ExportRenderError(
            "当前环境缺少 WeasyPrint 的系统依赖,无法导出 PDF。请安装 pango/gobject/cairo 后重试。"
        ) from exc
    return HTML(string=document_html).write_pdf()


def plan_to_docx(plan: VideoPlan) -> bytes:
    doc = Document()
    doc.add_heading(plan.title or "未命名方案", level=0)
    doc.add_paragraph(plan.summary or "暂无简介")

    doc.add_heading("基本信息", level=1)
    info = [
        ("方向", plan.direction),
        ("分类", plan.get_category_display()),
        ("AI 生成视频", "是" if plan.is_ai_generated_video else "否"),
        ("目标平台", plan.target_platform or "-"),
        ("目标观众", plan.target_audience or "-"),
        ("时长 (秒)", plan.duration_seconds),
        ("风格", plan.style or "-"),
        ("状态", plan.get_status_display()),
        ("最后更新", plan.updated_at.strftime("%Y-%m-%d %H:%M")),
    ]
    _add_key_values(doc, info)

    _add_mapping_section(doc, "内容设定", plan.content or {})
    _add_storyboard(doc, plan.storyboard or [])
    _add_mapping_section(doc, "剪辑建议", plan.editing_advice or {})
    _add_mapping_section(doc, "AI Prompt", plan.ai_prompts or {})

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def series_to_docx(series: SeriesPlan) -> bytes:
    doc = Document()
    doc.add_heading(series.title or "未命名系列", level=0)
    doc.add_paragraph(series.summary or "暂无简介")

    doc.add_heading("基本信息", level=1)
    _add_key_values(
        doc,
        [
            ("方向", series.direction),
            ("目标平台", series.target_platform or "-"),
            ("目标观众", series.target_audience or "-"),
            ("更新频率", series.update_frequency or "-"),
            ("单集时长 (秒)", series.episode_duration_seconds),
            ("计划集数", series.planned_episodes),
            ("状态", series.get_status_display()),
            ("最后更新", series.updated_at.strftime("%Y-%m-%d %H:%M")),
        ],
    )

    _add_mapping_section(doc, "系列定位", series.positioning or {})
    _add_mapping_section(doc, "单集模板", series.episode_template or {})
    _add_mapping_section(doc, "视觉风格", series.visual_style or {})
    _add_mapping_section(doc, "标题风格", series.title_style or {})
    _add_list_section(doc, "初始选题", series.initial_topics or [])
    _add_assets_sections(doc, series)
    _add_episodes_section(doc, series)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_key_values(doc: Document, pairs: list[tuple[str, Any]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "字段"
    header[1].text = "内容"
    for key, value in pairs:
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = _stringify(value)


def _add_mapping_section(doc: Document, title: str, data: dict[str, Any]) -> None:
    if not data:
        return
    doc.add_heading(title, level=1)
    for key, value in data.items():
        para = doc.add_paragraph()
        para.add_run(f"{key}: ").bold = True
        para.add_run(_stringify(value))


def _add_list_section(doc: Document, title: str, items: list[Any]) -> None:
    if not items:
        return
    doc.add_heading(title, level=1)
    for item in items:
        doc.add_paragraph(_stringify(item), style="List Number")


def _add_storyboard(doc: Document, shots: list[dict[str, Any]]) -> None:
    if not shots:
        return
    doc.add_heading("分镜脚本", level=1)
    headers = ["#", "时长", "画面 / 场景", "台词 / 旁白", "剪辑", "AI Prompt"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, title in enumerate(headers):
        table.rows[0].cells[index].text = title
    for i, shot in enumerate(shots, 1):
        cells = table.add_row().cells
        values = [
            i,
            _first_value(shot, ("duration", "seconds")),
            _first_value(shot, ("visual", "scene", "description")),
            _first_value(shot, ("line", "voiceover", "dialogue", "narration")),
            _first_value(shot, ("editing", "camera", "shot")),
            _first_value(shot, ("ai_prompt", "prompt")),
        ]
        for index, value in enumerate(values):
            cells[index].text = _stringify(value)


def _add_assets_sections(doc: Document, series: SeriesPlan) -> None:
    groups = [
        ("人物资产", series.characters.all()),
        ("风格资产", series.styles.all()),
        ("世界观资产", series.worldviews.all()),
        ("栏目资产", series.columns.all()),
    ]
    for title, queryset in groups:
        assets = list(queryset)
        if not assets:
            continue
        doc.add_heading(title, level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        headers = table.rows[0].cells
        headers[0].text = "名称"
        headers[1].text = "固定特征"
        headers[2].text = "详细设定"
        for asset in assets:
            cells = table.add_row().cells
            cells[0].text = asset.name
            cells[1].text = _stringify(asset.fixed_traits)
            cells[2].text = _stringify(asset.payload)


def _add_episodes_section(doc: Document, series: SeriesPlan) -> None:
    episodes = list(series.episodes.all())
    if not episodes:
        return
    doc.add_heading("单集清单", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["#", "标题", "状态", "时长", "更新时间"]
    for index, title in enumerate(headers):
        table.rows[0].cells[index].text = title
    for index, episode in enumerate(episodes, 1):
        cells = table.add_row().cells
        cells[0].text = str(index)
        cells[1].text = episode.title or "未命名单集"
        cells[2].text = episode.get_status_display()
        cells[3].text = f"{episode.duration_seconds}s"
        cells[4].text = episode.updated_at.strftime("%Y-%m-%d %H:%M")


def _first_value(data: dict[str, Any], keys: tuple[str, ...]) -> Any:
    for key in keys:
        value = data.get(key)
        if value not in (None, ""):
            return value
    return ""


def _stringify(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False, indent=2)
    return html.unescape(str(value))
